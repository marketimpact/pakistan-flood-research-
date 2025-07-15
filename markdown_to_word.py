#!/usr/bin/env python3
"""
Convert Markdown report to Word document format
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from datetime import datetime

def create_word_document(markdown_file, output_file):
    """Convert markdown to Word document with proper formatting"""
    
    # Create new document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split content into lines
    lines = content.split('\n')
    
    # Process each line
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # Empty line - add paragraph break
            doc.add_paragraph()
            i += 1
            continue
            
        if line.startswith('# '):
            # Main heading (H1)
            heading = line[2:].strip()
            h = doc.add_heading(heading, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line.startswith('## '):
            # Section heading (H2)
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
            
        elif line.startswith('### '):
            # Subsection heading (H3)
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)
            
        elif line.startswith('#### '):
            # Sub-subsection heading (H4)
            heading = line[5:].strip()
            doc.add_heading(heading, level=4)
            
        elif line.startswith('**') and line.endswith('**'):
            # Bold standalone line
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point - collect all consecutive bullets
            bullets = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                bullet_text = lines[i].strip()[2:].strip()
                bullets.append(bullet_text)
                i += 1
            
            # Add all bullets
            for bullet in bullets:
                p = doc.add_paragraph(bullet, style='List Bullet')
            
            # Continue without incrementing i (already done in loop)
            continue
            
        elif line.startswith('1. ') or re.match(r'^\d+\. ', line):
            # Numbered list - collect all consecutive numbers
            numbers = []
            while i < len(lines) and re.match(r'^\d+\. ', lines[i].strip()):
                number_text = re.sub(r'^\d+\. ', '', lines[i].strip())
                numbers.append(number_text)
                i += 1
            
            # Add all numbered items
            for number in numbers:
                p = doc.add_paragraph(number, style='List Number')
            
            # Continue without incrementing i
            continue
            
        elif line.startswith('|') and '|' in line[1:]:
            # Table - collect all table rows
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if '---' not in lines[i]:  # Skip separator rows
                    row_data = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                    table_rows.append(row_data)
                i += 1
            
            if table_rows:
                # Create table
                table = doc.add_table(rows=1, cols=len(table_rows[0]))
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                for j, cell_text in enumerate(table_rows[0]):
                    header_cells[j].text = cell_text
                    # Make header bold
                    for paragraph in header_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for row_data in table_rows[1:]:
                    row_cells = table.add_row().cells
                    for j, cell_text in enumerate(row_data):
                        if j < len(row_cells):
                            row_cells[j].text = cell_text
            
            # Continue without incrementing i
            continue
            
        elif line.startswith('---'):
            # Horizontal rule - add page break or line
            doc.add_page_break()
            
        elif line.startswith('```'):
            # Code block - collect until closing ```
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Add code block
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Intense Quote'
            
        else:
            # Regular paragraph
            if line:
                # Process inline formatting
                p = doc.add_paragraph()
                
                # Split by ** for bold formatting
                parts = line.split('**')
                for j, part in enumerate(parts):
                    if j % 2 == 0:
                        # Regular text
                        if part:
                            p.add_run(part)
                    else:
                        # Bold text
                        if part:
                            run = p.add_run(part)
                            run.bold = True
        
        i += 1
    
    # Add footer with generation info
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated on {datetime.now().strftime('%B %d, %Y')} | Pakistan Virtual Gauge Analysis Report"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    doc.save(output_file)
    print(f"Word document saved as: {output_file}")

def main():
    """Convert the Pakistan report to Word format"""
    
    markdown_file = "Pakistan_Virtual_Gauge_Analysis_Report.md"
    output_file = "Pakistan_Virtual_Gauge_Analysis_Report.docx"
    
    try:
        create_word_document(markdown_file, output_file)
        print(f"✅ Successfully converted {markdown_file} to {output_file}")
        print(f"📄 Document contains comprehensive analysis of Pakistan's virtual gauge network")
        print(f"📊 Includes 2,391 gauge analysis with technical details and recommendations")
        
    except Exception as e:
        print(f"❌ Error converting to Word: {e}")

if __name__ == "__main__":
    main()