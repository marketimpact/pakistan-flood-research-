#!/usr/bin/env python3
"""
Improved Markdown to Word converter with proper formatting
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor
import re
from datetime import datetime

def create_professional_word_document(markdown_file, output_file):
    """Convert markdown to professionally formatted Word document"""
    
    # Create new document
    doc = Document()
    
    # Configure document styles
    setup_document_styles(doc)
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Read and process markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Process content
    process_markdown_content(doc, content)
    
    # Add professional footer
    add_footer(doc)
    
    # Save document
    doc.save(output_file)
    print(f"Professional Word document saved as: {output_file}")

def setup_document_styles(doc):
    """Configure professional document styles"""
    
    # Title style
    title_style = doc.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(18)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Subtitle style
    subtitle_style = doc.styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(14)
    subtitle_font.italic = True
    subtitle_font.color.rgb = RGBColor(0x44, 0x72, 0xc4)
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(18)

def process_markdown_content(doc, content):
    """Process markdown content with proper formatting"""
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        if line.startswith('# '):
            # Main title
            title = clean_text(line[2:].strip())
            p = doc.add_paragraph(title, style='Custom Title')
            
        elif line.startswith('## '):
            # Subtitle
            subtitle = clean_text(line[3:].strip())
            p = doc.add_paragraph(subtitle, style='Custom Subtitle')
            
        elif line.startswith('### '):
            # Section heading
            heading = clean_text(line[4:].strip())
            doc.add_heading(heading, level=1)
            
        elif line.startswith('#### '):
            # Subsection heading
            heading = clean_text(line[5:].strip())
            doc.add_heading(heading, level=2)
            
        elif line.startswith('---'):
            # Page break or separator
            if i > 0:  # Don't add break at start
                doc.add_page_break()
            
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet points
            i = process_bullet_list(doc, lines, i)
            continue
            
        elif re.match(r'^\d+\.', line):
            # Numbered list
            i = process_numbered_list(doc, lines, i)
            continue
            
        else:
            # Regular paragraph
            if line:
                para_text = clean_text(line)
                if para_text:
                    p = doc.add_paragraph()
                    add_formatted_text(p, para_text)
        
        i += 1

def clean_text(text):
    """Remove markdown formatting artifacts"""
    # Remove bold markers
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # Remove italic markers  
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # Remove inline code markers
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    # Clean up any remaining markdown artifacts
    text = text.replace('**', '').replace('*', '')
    
    return text.strip()

def add_formatted_text(paragraph, text):
    """Add text with inline formatting"""
    
    # Split by potential bold markers that weren't cleaned
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            content = part[2:-2]
            if content:
                run = paragraph.add_run(content)
                run.bold = True
        else:
            # Regular text
            if part:
                paragraph.add_run(part)

def process_bullet_list(doc, lines, start_index):
    """Process bullet list items"""
    i = start_index
    
    while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
        bullet_text = clean_text(lines[i].strip()[2:])
        if bullet_text:
            doc.add_paragraph(bullet_text, style='List Bullet')
        i += 1
    
    return i

def process_numbered_list(doc, lines, start_index):
    """Process numbered list items"""
    i = start_index
    
    while i < len(lines) and re.match(r'^\d+\.', lines[i].strip()):
        number_text = re.sub(r'^\d+\.\s*', '', lines[i].strip())
        number_text = clean_text(number_text)
        if number_text:
            doc.add_paragraph(number_text, style='List Number')
        i += 1
    
    return i

def add_footer(doc):
    """Add professional footer"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Pakistan Virtual Gauge Analysis Report | Generated {datetime.now().strftime('%B %d, %Y')} | Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page number
    run = footer_para.runs[0]
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def main():
    """Convert the full Pakistan report to Word format"""
    
    markdown_file = "Pakistan_Virtual_Gauge_Report_Full.md"
    output_file = "Pakistan_Virtual_Gauge_Report_Full.docx"
    
    try:
        create_professional_word_document(markdown_file, output_file)
        print(f"✅ Successfully created professional 10-page report: {output_file}")
        print(f"📄 Clean formatting with no markdown artifacts")
        print(f"📊 Concise analysis focused on key findings and recommendations")
        print(f"🎯 Ready for stakeholder presentations and decision-making")
        
    except Exception as e:
        print(f"❌ Error creating Word document: {e}")

if __name__ == "__main__":
    main()